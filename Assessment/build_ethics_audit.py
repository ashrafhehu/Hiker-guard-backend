import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import win32com.client

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def prevent_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

def make_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def build_ethics_audit():
    input_file = "CERTIFIED ASSIGNMENT 2_MODULE 1 SESSION 2 GeoAI Hackathon Session1_Intro_Ethics_v1.docx"
    output_docx = "AGAIF2026_BC1_CA2_MY-411_Muhammad Ashraf.docx"
    output_pdf = "AGAIF2026_BC1_CA2_MY-411_Muhammad Ashraf.pdf"

    doc = Document(input_file)

    # 1. Update Existing Paragraphs at the top (P0 to P3)
    p0 = doc.paragraphs[0]
    p0.text = "CERTIFIED ASSESSMENT ASSIGNMENT 2: RESPONSIBLE AI ETHICS AUDIT"
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.runs[0].font.name = "Arial"
    p0.runs[0].font.size = Pt(14)
    p0.runs[0].font.bold = True
    p0.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    p1 = doc.paragraphs[1]
    p1.text = "Evaluation of AI-Based Geospatial Flood Early Warning System"
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].font.name = "Arial"
    p1.runs[0].font.size = Pt(11)
    p1.runs[0].font.italic = True
    p1.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    # Insert Metadata Paragraph after P1
    p_meta = p1.insert_paragraph_before()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.space_before = Pt(6)
    p_meta.paragraph_format.space_after = Pt(10)

    meta_items = [
        ("Participant Name: ", "Muhammad Ashraf"),
        ("AGAIF Reference Code: ", "MY-411"),
        ("Module / Task: ", "Module 1 Session 2 – GeoAI Ethics Audit Exercise (CA2)"),
        ("Target System Evaluated: ", "AI-Based Geospatial Flood Warning System"),
        ("Assessment Date: ", "August 11, 2026")
    ]

    for label, val in meta_items:
        r_lbl = p_meta.add_run(label)
        r_lbl.font.name = "Arial"
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor(30, 58, 138)

        r_val = p_meta.add_run(val + ("  |  " if label != meta_items[-1][0] else ""))
        r_val.font.name = "Arial"
        r_val.font.size = Pt(9.5)
        r_val.font.bold = False
        r_val.font.color.rgb = RGBColor(30, 41, 59)

    # Format remaining instructions paragraphs nicely
    for i in range(2, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        if p.text.strip():
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(9)

    # 2. Populate and Style Audit Table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Table Header Styling
        hdr_row = table.rows[0]
        make_row_header(hdr_row)
        prevent_row_split(hdr_row)
        hdr_cells = hdr_row.cells
        hdr_titles = ["No.", "Dimension", "Green\n(Acceptable)", "Yellow\n(Improve)", "Red\n(High Risk)", "Justification, Identified Ethical Risks, Stakeholder Impact & Mitigation Measures"]
        
        for idx, cell in enumerate(hdr_cells):
            cell.text = hdr_titles[idx]
            set_cell_background(cell, "1E3A8A") # Navy header #1e3a8a
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

        # Dimension Audit Rows Data
        audit_rows_data = [
            {
                "no": "1.",
                "dim": "Fairness",
                "rating": "Yellow",
                "notes": [
                    ("Current Condition & Justification: ", "Model accuracy is predominantly tested in heavily gauged urban river basins (e.g., Klang Valley). Rural, indigenous (Orang Asli), and informal settlement river catchments suffer from sparse gauge density, resulting in spatial data under-representation."),
                    ("Identified Ethical Risk: ", "Algorithmic Bias & Spatial Disparity — Higher false-negative rate (missed flood alerts) in low-gauge rural zones."),
                    ("Impact & Stakeholders: ", "Rural agricultural communities and low-income informal settlers face unannounced flood inundation, leading to severe property destruction, livestock loss, and potential casualties."),
                    ("Mitigation Measures: ", "1) Integrate Sentinel-1 SAR satellite water masks to supplement low-gauge rural basins. 2) Deploy low-cost IoT stream gauges in remote hot-spots. 3) Re-weight training loss to equalize spatial prediction sensitivity across urban and rural catchments.")
                ]
            },
            {
                "no": "2.",
                "dim": "Privacy",
                "rating": "Red",
                "notes": [
                    ("Current Condition & Justification: ", "The system ingests and centrally stores exact household GPS coordinates, personal phone numbers, and residential address records paired with flood-vulnerability scores without anonymization or encryption."),
                    ("Identified Ethical Risk: ", "Privacy Infringement & Unsanctioned Surveillance — Central storage of raw, unencrypted PII exposes sensitive residential location data."),
                    ("Impact & Stakeholders: ", "Property owners and residents risk data leakage, unauthorized commercial exploitation, real-estate devaluation, or predatory insurance rate discrimination."),
                    ("Mitigation Measures: ", "1) Enforce spatial k-anonymity and aggregate public alert dispatch to 250m H3 hexagonal spatial grids rather than raw household points. 2) Implement end-to-end PII encryption and Role-Based Access Control (RBAC). 3) Enforce strict data minimization and purge transient contact telemetry post-alert delivery.")
                ]
            },
            {
                "no": "3.",
                "dim": "Transparency",
                "rating": "Yellow",
                "notes": [
                    ("Current Condition & Justification: ", "The AI system outputs opaque, black-box probability scores (e.g., 'Risk = 0.88') without explaining the underlying spatial drivers or physical triggers."),
                    ("Identified Ethical Risk: ", "Opacity & Explainability Deficit — Emergency managers and residents receive unexplained numbers without physical context."),
                    ("Impact & Stakeholders: ", "District emergency officers (BOMBA, SMART) experience decision hesitation during evacuations, while residents lack clarity on flood severity drivers (e.g., dam release vs. extreme rainfall)."),
                    ("Mitigation Measures: ", "1) Integrate Explainable AI (XAI) feature attribution maps (SHAP/Integrated Gradients) into emergency dashboards. 2) Convert probabilities into natural-language alerts (e.g., 'High risk driven by 120mm/hr upstream rainfall combined with 16:00 high tide'). 3) Publish an open Model Card detailing dataset boundaries, accuracy limitations, and operating constraints.")
                ]
            },
            {
                "no": "4.",
                "dim": "Accountability",
                "rating": "Red",
                "notes": [
                    ("Current Condition & Justification: ", "No formal governance framework or legal RACI matrix exists defining explicit liability between GeoAI developers, hydrologists, and public emergency response agencies in the event of system failure."),
                    ("Identified Ethical Risk: ", "Diffusion of Responsibility & Unclear Liability — Absence of a designated entity held accountable for false positives (costly panic) or false negatives (unwarned disasters)."),
                    ("Impact & Stakeholders: ", "Disaster response agencies face legal vulnerabilities and operational deadlock, while affected flood victims are left without legal recourse or formal redress mechanisms."),
                    ("Mitigation Measures: ", "1) Formulate an Inter-Agency Memorandum of Understanding (MoU) establishing that the AI operates strictly as decision support, while human Civil Defense Directors retain final decision authority. 2) Mandate an independent Post-Event AI Performance Audit Committee. 3) Maintain an immutable, tamper-evident audit log of all predictions and human overrides.")
                ]
            },
            {
                "no": "5.",
                "dim": "Safety & Reliability",
                "rating": "Yellow",
                "notes": [
                    ("Current Condition & Justification: ", "The model lacks real-time confidence metrics, Out-of-Distribution (OOD) monitoring, and automated fallback controls when input sensors (telemetry/radar) become corrupted or degraded during extreme storms."),
                    ("Identified Ethical Risk: ", "System Degradation & Misleading Precision — Risk of confident but dangerously incorrect predictions during unprecedented climate events exceeding historical training data."),
                    ("Impact & Stakeholders: ", "Emergency responders and vulnerable populations risk relying on degraded AI outputs, resulting in misallocated rescue boats, missed evacuations, or fatal delays."),
                    ("Mitigation Measures: ", "1) Deploy real-time OOD distance monitors and sensor data sanity checks. 2) Display explicit spatial confidence bands (High/Medium/Low) on responder maps. 3) Implement automated fail-safe fallback to deterministic physics-based hydrological models (HEC-RAS/SWMM) when ML uncertainty exceeds safety thresholds.")
                ]
            },
            {
                "no": "6.",
                "dim": "Sustainability",
                "rating": "Green",
                "notes": [
                    ("Current Condition & Justification: ", "The system utilizes lightweight, optimized spatial inference pipelines (pruned LightGBM/U-Net) running on modest, energy-efficient cloud infrastructure with minimal carbon footprint."),
                    ("Identified Ethical Risk: ", "Low Environmental Risk — Operational energy usage and computational costs are well-managed and sustainable."),
                    ("Impact & Stakeholders: ", "Positive long-term financial and environmental viability for municipal disaster budgets."),
                    ("Mitigation Measures: ", "1) Maintain green computing standards by hosting cloud retraining pipelines in regions powered by renewable energy. 2) Utilize incremental delta-training updates rather than full retraining runs to minimize compute overhead.")
                ]
            },
            {
                "no": "7.",
                "dim": "Human Oversight",
                "rating": "Green",
                "notes": [
                    ("Current Condition & Justification: ", "Designed with a mandatory Human-in-the-Loop (HITL) protocol where automated AI alerts require dual sign-off from duty hydrologists before public broadcast, supported by community crowd-sourced ground-truth reporting."),
                    ("Identified Ethical Risk: ", "Adequately Controlled Risk — Human verification effectively prevents unchecked automated errors."),
                    ("Impact & Stakeholders: ", "Protects the public from false panics while empowering local residents to submit real-time ground-truth flood reports."),
                    ("Mitigation Measures: ", "1) Conduct regular scenario training for duty officers to prevent 'automation bias' (blindly trusting AI) or 'alert fatigue'. 2) Maintain community ground-truth reporting loops to validate and calibrate AI prediction grids in real time.")
                ]
            }
        ]

        # Populate rows
        for r_idx, data in enumerate(audit_rows_data):
            row_obj = table.rows[r_idx + 1]
            prevent_row_split(row_obj)
            row_cells = row_obj.cells
            
            # Column 0: No.
            row_cells[0].text = data["no"]
            p0 = row_cells[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p0.runs[0].font.bold = True
            p0.runs[0].font.size = Pt(9)
            
            # Column 1: Dimension
            row_cells[1].text = data["dim"]
            p1 = row_cells[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p1.runs[0].font.bold = True
            p1.runs[0].font.size = Pt(9.5)
            p1.runs[0].font.color.rgb = RGBColor(30, 58, 138)
            
            # Columns 2, 3, 4: Green, Yellow, Red traffic lights
            r_type = data["rating"]
            row_cells[2].text = "✔ ACCEPTABLE" if r_type == "Green" else ""
            row_cells[3].text = "⚠ IMPROVE" if r_type == "Yellow" else ""
            row_cells[4].text = "✖ HIGH RISK" if r_type == "Red" else ""

            # Highlight rating cell background
            if r_type == "Green":
                set_cell_background(row_cells[2], "DCFCE7") # Light Green #dcfce7
                p_g = row_cells[2].paragraphs[0]
                p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p_g.runs:
                    p_g.runs[0].font.bold = True
                    p_g.runs[0].font.size = Pt(8.5)
                    p_g.runs[0].font.color.rgb = RGBColor(22, 101, 52)
            elif r_type == "Yellow":
                set_cell_background(row_cells[3], "FEF08A") # Light Yellow #fef08a
                p_y = row_cells[3].paragraphs[0]
                p_y.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p_y.runs:
                    p_y.runs[0].font.bold = True
                    p_y.runs[0].font.size = Pt(8.5)
                    p_y.runs[0].font.color.rgb = RGBColor(133, 77, 14)
            elif r_type == "Red":
                set_cell_background(row_cells[4], "FECACA") # Light Red #fecaca
                p_r = row_cells[4].paragraphs[0]
                p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p_r.runs:
                    p_r.runs[0].font.bold = True
                    p_r.runs[0].font.size = Pt(8.5)
                    p_r.runs[0].font.color.rgb = RGBColor(153, 27, 27)

            # Column 5: Notes & Detailed Justifications
            cell_notes = row_cells[5]
            cell_notes.text = "" # Clear default
            set_cell_margins(cell_notes, top=100, bottom=100, left=120, right=120)

            for n_idx, (prefix, text) in enumerate(data["notes"]):
                p_n = cell_notes.paragraphs[0] if n_idx == 0 else cell_notes.add_paragraph()
                p_n.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_n.paragraph_format.space_after = Pt(3)
                p_n.paragraph_format.line_spacing = 1.12

                r_pre = p_n.add_run(prefix)
                r_pre.font.name = "Arial"
                r_pre.font.bold = True
                r_pre.font.size = Pt(8.5)
                r_pre.font.color.rgb = RGBColor(30, 58, 138)

                r_txt = p_n.add_run(text)
                r_txt.font.name = "Arial"
                r_txt.font.bold = False
                r_txt.font.size = Pt(8.5)
                r_txt.font.color.rgb = RGBColor(30, 41, 59)

    # 3. Overall Deployment Recommendation Section
    doc.add_page_break() # Clean page for executive recommendation

    p_rec_head = doc.add_paragraph()
    p_rec_head.paragraph_format.space_before = Pt(12)
    p_rec_head.paragraph_format.space_after = Pt(8)
    r_rh = p_rec_head.add_run("OVERALL DEPLOYMENT RECOMMENDATION & CONCLUSION")
    r_rh.font.name = "Arial"
    r_rh.font.size = Pt(14)
    r_rh.font.bold = True
    r_rh.font.color.rgb = RGBColor(15, 23, 42)

    # Recommendation Box / Highlight
    p_verdict = doc.add_paragraph()
    p_verdict.paragraph_format.space_after = Pt(10)
    r_v_lbl = p_verdict.add_run("FINAL AUDIT VERDICT: ")
    r_v_lbl.font.name = "Arial"
    r_v_lbl.font.size = Pt(12)
    r_v_lbl.font.bold = True
    r_v_lbl.font.color.rgb = RGBColor(15, 23, 42)

    r_v_val = p_verdict.add_run("PROCEED SUBJECT TO SPECIFIED CONDITIONS (Conditioned Approval)")
    r_v_val.font.name = "Arial"
    r_v_val.font.size = Pt(12)
    r_v_val.font.bold = True
    r_v_val.font.color.rgb = RGBColor(180, 83, 9) # Amber / Dark Yellow #b45309

    # Detailed Justification Paragraphs
    p_exp = doc.add_paragraph()
    p_exp.paragraph_format.space_after = Pt(8)
    p_exp.paragraph_format.line_spacing = 1.15
    r_exp = p_exp.add_run(
        "Executive Summary & Decision Rationale:\n"
        "The proposed AI-Based Flood Early Warning System exhibits significant potential value for public safety, disaster preparedness, and rapid emergency response. "
        "However, full regional deployment must NOT proceed in its current un-mitigated state due to two (2) critical RED high-risk dimensions (Privacy and Accountability) "
        "and three (3) YELLOW dimensions (Fairness, Transparency, and Safety & Reliability)."
    )
    r_exp.font.name = "Arial"
    r_exp.font.size = Pt(10)
    r_exp.font.color.rgb = RGBColor(51, 65, 85)

    p_cond_head = doc.add_paragraph()
    p_cond_head.paragraph_format.space_after = Pt(4)
    r_ch = p_cond_head.add_run("Mandatory Pre-Deployment Conditions (Prerequisites for Pilot Launch):")
    r_ch.font.name = "Arial"
    r_ch.font.size = Pt(10.5)
    r_ch.font.bold = True
    r_ch.font.color.rgb = RGBColor(30, 58, 138)

    conditions = [
        ("1. Privacy Remediation (RED Risk): ", "Eliminate central storage of raw household GPS coordinates and unencrypted PII. Enforce spatial k-anonymity (250m H3 hex grid aggregation) and complete an independent Privacy Impact Assessment (PIA)."),
        ("2. Legal Accountability Framework (RED Risk): ", "Ratify an Inter-Agency Memorandum of Understanding (MoU) formally declaring that the AI operates strictly as decision support, while human Civil Defense Directors retain sole legal evacuation authority. Establish an immutable audit logging pipeline."),
        ("3. Safety & Fallback Controls (YELLOW Risk): ", "Implement real-time sensor sanity monitors and an automated fail-safe fallback to physics-based hydrological models (HEC-RAS) when ML uncertainty is high."),
        ("4. Spatial Fairness Integration (YELLOW Risk): ", "Supplement sparse rural stream gauge telemetries with Sentinel-1 SAR satellite water masks and low-cost IoT gauges to eliminate urban-rural accuracy disparities."),
        ("5. Phased Pilot Rollout (90-Day Controlled Pilot): ", "Deploy the system exclusively as a limited 90-day pilot within a single controlled river basin (e.g., Selangor River catchment) under active human-in-the-loop validation before national expansion.")
    ]

    for cond_title, cond_desc in conditions:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(4)
        p_c.paragraph_format.line_spacing = 1.12

        r_ct = p_c.add_run("• " + cond_title)
        r_ct.font.name = "Arial"
        r_ct.font.bold = True
        r_ct.font.size = Pt(9.5)
        r_ct.font.color.rgb = RGBColor(30, 58, 138)

        r_cd = p_c.add_run(cond_desc)
        r_cd.font.name = "Arial"
        r_cd.font.bold = False
        r_cd.font.size = Pt(9.5)
        r_cd.font.color.rgb = RGBColor(30, 41, 59)

    # Sign-off Box
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(16)
    r_sign = p_sign.add_run(
        "Ethics Audit Conducted By:\n"
        "Muhammad Ashraf (Participant Ref: MY-411)\n"
        "Lead GeoAI Ethics Auditor & Systems Reviewer\n"
        "ASEAN GeoAI Fusion 2026 Hackathon"
    )
    r_sign.font.name = "Arial"
    r_sign.font.size = Pt(9.5)
    r_sign.font.italic = True
    r_sign.font.color.rgb = RGBColor(71, 85, 105)

    # Save DOCX
    doc.save(output_docx)
    print(f"Successfully generated DOCX: {output_docx}")

    # Convert DOCX to PDF using Word COM
    abs_docx = os.path.abspath(output_docx)
    abs_pdf = os.path.abspath(output_pdf)

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_com = word.Documents.Open(abs_docx)
    doc_com.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
    doc_com.Close()
    word.Quit()
    print(f"Successfully generated PDF: {output_pdf} (Size: {os.path.getsize(abs_pdf)} bytes)")

if __name__ == "__main__":
    build_ethics_audit()
